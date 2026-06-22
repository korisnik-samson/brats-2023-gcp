"""
Assemble the master thesis into a .docx that matches the candidate's BACHELOR thesis exactly.

Strategy: open the bachelor thesis as the BASE document (inheriting its styles, the
style-linked heading auto-numbering, the "- X -" footer, A4 / 3 cm left margin, Title /
Subtitle / Body Text / Bibliography / TOC styles), clear its body, and repopulate with the
master-thesis content using those same styles. This guarantees identical formatting,
arrangement and alignment.

  python build_docx.py   ->  thesis/Master_Thesis_Offorjindu.docx

Notes
- Headings are AUTO-numbered by the inherited Heading 1/2/3 styles, so the manual "1.", "1.1"
  numbers in the Markdown are stripped. Abstract and Literature are unnumbered (numId 0).
- Citations follow the bachelor's inline author-date style (already how the chapters are
  written); the Literature list uses the Bibliography style. If the candidate wants live
  Zotero fields, they can re-insert citations in Word — the text format already matches.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

THESIS = Path("thesis")
BASE = THESIS / "Samson Offorjindu - Diplomski rad.docx"     # bachelor = formatting template
OUT = THESIS / "Master_Thesis_Offorjindu.docx"

TITLE = ("ONE PIPELINE, THREE TUMOURS: UNIFIED DEEP-LEARNING SEGMENTATION OF ADULT, "
         "MENINGIOMA AND PAEDIATRIC BRAIN TUMOURS ON A SINGLE GPU (BRATS 2023)")
MENTOR = "Professor Nebojša Bačanin Džakula, PhD"
CANDIDATE = "Samson O. Offorjindu"
YEAR = "2026."

CHAPTERS = ["01_introduction.md", "02_background.md", "03_methodology.md",
            "04_results.md", "05_discussion.md", "06_conclusion.md"]

STATUS_RE = re.compile(r"\s*(✅|⏳|◻|⭐)+.*$")
TAG_RE = re.compile(r"\s*\[(verify|fill|pending|update[^\]]*|confirm[^\]]*)\]", re.I)
INLINE_TAG = re.compile(r"\s*\*{0,2}\[(?:verify|fill|pending|update[^\]]*|confirm[^\]]*)\]\*{0,2}", re.I)
NUM_PREFIX = re.compile(r"^\d+(\.\d+)*\.?\s+")           # "1. ", "4.5 ", "3.10 "
TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
SKIP_HINTS = ("drafting note", "to finalise", "to finalize", "still to add", "status.")


def clean_heading(t):
    t = STATUS_RE.sub("", t)
    t = TAG_RE.sub("", t)
    t = re.sub(r"[*`]", "", t)
    t = re.sub(r"\s*[—\-]\s*za proveru.*$", "", t, flags=re.I)
    t = NUM_PREFIX.sub("", t)                            # drop manual heading numbers
    return t.strip()


def add_runs(p, text, bold=False, italic=False):
    text = INLINE_TAG.sub("", text).replace("->", "→")
    for part in TOKEN.split(text):
        if not part:
            continue
        b, i = bold, italic
        if part.startswith("**") and part.endswith("**"):
            part, b = part[2:-2], True
        elif part.startswith("*") and part.endswith("*"):
            part, i = part[1:-1], True
        elif part.startswith("`") and part.endswith("`"):
            part = part[1:-1]
        r = p.add_run(part); r.bold = b; r.italic = i


def clear_body(doc):
    """Remove all body content but keep the trailing sectPr (page setup + footer)."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def field(p, instr):
    for kind, txt in [("begin", None), (None, instr), ("separate", None), ("end", None)]:
        r = p.add_run()
        if txt is not None:
            e = OxmlElement("w:instrText"); e.set(qn("xml:space"), "preserve"); e.text = txt
        else:
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), kind)
        r._r.append(e)


def suppress_number(p):
    """Remove this heading from the auto-number list (numId 0) -> unnumbered."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    for tag, val in [("w:ilvl", "0"), ("w:numId", "0")]:
        e = OxmlElement(tag); e.set(qn("w:val"), val); numPr.append(e)
    pPr.append(numPr)


def para(doc, text="", style=None, align=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        add_runs(p, text, bold=bold, italic=italic)
    if size:
        for r in p.runs:
            r.font.size = Pt(size)
    return p


def heading(doc, text, level, numbered=True):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if not numbered:
        suppress_number(p)
    return p


# ── title page (matches the bachelor layout) ──────────────────────────────────
def title_page(doc):
    C = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "SINGIDUNUM UNIVERSITY", align=C, bold=True)
    para(doc, "FACULTY OF TECHNICAL SCIENCES", align=C, bold=True)
    for _ in range(3):
        para(doc)
    para(doc, TITLE, style="Title")
    para(doc, "- Master Thesis -", style="Subtitle")
    for _ in range(4):
        para(doc)
    t = doc.add_table(rows=2, cols=2)
    t.columns[0].width = t.columns[1].width
    t.cell(0, 0).paragraphs[0].add_run("Mentor:")
    t.cell(0, 1).paragraphs[0].add_run("Candidate:")
    t.cell(1, 0).paragraphs[0].add_run(MENTOR)
    t.cell(1, 1).paragraphs[0].add_run(CANDIDATE)
    for _ in range(6):
        para(doc)
    para(doc, f"Belgrade, {YEAR}", align=C)
    doc.add_page_break()


def contents(doc):
    try:
        doc.add_paragraph("Contents", style="TOC Heading")   # not itself in the TOC
    except KeyError:
        heading(doc, "Contents", 1, numbered=False)
    p = doc.add_paragraph()
    field(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


def _table_borders(tbl):
    """Apply single-line borders to all cells (no named style needed)."""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto")
        borders.append(e)
    tblPr.append(borders)


# ── markdown table -> Word table ──────────────────────────────────────────────
def render_table(doc, content):
    rows, notes, caption = [], [], None
    for c in content:
        if c.startswith("|"):
            if re.match(r"^\|[\s:\-|]+\|?$", c):
                continue
            rows.append([x.strip() for x in c.strip("|").split("|")])
        elif c.strip():
            (notes if caption else [caption])  # noop
            if caption is None:
                caption = c
            else:
                notes.append(c)
    if caption:
        para(doc, NUM_PREFIX.sub("", clean_heading(caption)), bold=True, size=11)
    if rows:
        ncol = max(len(r) for r in rows)
        tbl = doc.add_table(rows=0, cols=ncol)
        _table_borders(tbl)
        for ri, cells in enumerate(rows):
            cs = tbl.add_row().cells
            for ci in range(ncol):
                cp = cs[ci].paragraphs[0]
                add_runs(cp, cells[ci] if ci < len(cells) else "", bold=(ri == 0))
                for r in cp.runs:
                    r.font.size = Pt(10)
    for n in notes:
        para(doc, n, italic=True, size=10)


# ── convert a chapter markdown file ───────────────────────────────────────────
def convert(doc, md, numbered=True):
    lines = Path(md).read_text(encoding="utf-8").splitlines()
    i, buf = 0, []

    def flush():
        nonlocal buf
        if buf:
            para(doc, " ".join(buf), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            buf = []

    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith(">"):
            flush()
            block = []
            while i < len(lines) and lines[i].startswith(">"):
                block.append(lines[i].lstrip(">").strip()); i += 1
            if any(h in " ".join(block).lower() for h in SKIP_HINTS):
                continue
            if any(c.startswith("|") for c in block):
                render_table(doc, block)
            continue
        if line.startswith("#"):
            flush()
            txt = clean_heading(line.lstrip("#").strip())
            if "references cited" in txt.lower():
                break
            lvl = min(len(line) - len(line.lstrip("#")), 3)
            heading(doc, txt, lvl, numbered=numbered)
            i += 1
            continue
        if re.match(r"^\*\(.+\)\*$", line):           # Serbian subtitle -> drop
            i += 1
            continue
        if line.strip() in ("", "---"):
            flush(); i += 1; continue
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            flush()
            item = m.group(3)
            while i + 1 < len(lines):
                nxt = lines[i + 1]
                if (not nxt.strip() or nxt.startswith("#") or nxt.startswith(">")
                        or re.match(r"^(\s*)([-*]|\d+\.)\s+", nxt)):
                    break
                item += " " + nxt.strip(); i += 1
            marker = m.group(2)
            prefix = (marker + " ") if marker[0].isdigit() else "• "
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(prefix)
            add_runs(p, item)
            i += 1
            continue
        buf.append(line.strip()); i += 1
    flush()


# ── abstract (English only, unnumbered "Abstract" heading) ────────────────────
def abstract(doc):
    heading(doc, "Abstract", 1, numbered=False)
    lines = (THESIS / "abstract.md").read_text(encoding="utf-8").splitlines()
    grab, buf = False, []

    def flush():
        if buf:
            para(doc, " ".join(buf), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            buf.clear()

    for ln in lines:
        s = ln.strip()
        if s.startswith("## Abstract (English)"):
            grab = True; continue
        if s.startswith("## Rezime"):
            break
        if not grab:
            continue
        if s in ("", "---") or s.startswith(">"):
            flush(); continue
        buf.append(s)
    flush()
    doc.add_page_break()


# ── Literature (Bibliography style, alphabetical, no numbers/bold) ─────────────
def literature(doc):
    heading(doc, "Literature", 1, numbered=False)
    for ln in (THESIS / "references.md").read_text(encoding="utf-8").splitlines():
        m = re.match(r"^\d+\.\s+(.*)$", ln.strip())
        if not m:
            continue
        text = m.group(1).replace("**", "")        # APA bibliography: no bold authors
        try:
            p = doc.add_paragraph(style="Bibliography")
        except KeyError:
            p = doc.add_paragraph()
        add_runs(p, INLINE_TAG.sub("", text))


def main():
    doc = Document(str(BASE))
    clear_body(doc)
    title_page(doc)
    contents(doc)
    abstract(doc)
    for ch in CHAPTERS:
        convert(doc, THESIS / ch, numbered=True)
        doc.add_page_break()
    literature(doc)
    doc.save(OUT)
    print(f"Saved {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    main()
